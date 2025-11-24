"""Output formatter node for creating final contract files."""

import os
from typing import Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.models.state import ContractState


def output_formatter_node(state: ContractState) -> Dict[str, Any]:
    """
    Format and save the final contract in multiple formats.
    """
    print("💾 Formatting and saving final contract...")

    updates = {
        "current_step": "output_formatter",
        "messages": []
    }

    contract_draft = state.get("contract_draft", "")
    if not contract_draft:
        updates["error"] = "No contract draft available for formatting"
        updates["messages"].append({
            "role": "system",
            "content": "❌ No contract draft available to format"
        })
        return updates

    try:
        # Create output directory if it doesn't exist
        output_dir = "data/output"
        os.makedirs(output_dir, exist_ok=True)

        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"contract_{timestamp}"

        # Save as plain text
        txt_path = os.path.join(output_dir, f"{base_filename}.txt")
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(contract_draft)

        updates["messages"].append({
            "role": "system",
            "content": f"✓ Saved text version: {txt_path}"
        })

        # Save as DOCX
        docx_path = os.path.join(output_dir, f"{base_filename}.docx")
        create_docx_contract(contract_draft, docx_path, state.get("merged_data", {}))

        updates["messages"].append({
            "role": "system",
            "content": f"✓ Saved DOCX version: {docx_path}"
        })

        # Save formatted version for display
        updates["formatted_contract"] = contract_draft
        updates["output_path"] = docx_path

        # Create summary
        lines = contract_draft.split('\n')
        word_count = len(contract_draft.split())

        updates["messages"].append({
            "role": "system",
            "content": f"""
✅ Contract generation complete!

📊 Statistics:
  • Lines: {len(lines)}
  • Words: {word_count}
  • Characters: {len(contract_draft)}

📁 Output files:
  • Text: {txt_path}
  • Word: {docx_path}

Quality Score: {state.get('quality_report', {}).get('score', 'N/A'):.1f}%
"""
        })

    except Exception as e:
        updates["error"] = f"Output formatting failed: {str(e)}"
        updates["messages"].append({
            "role": "system",
            "content": f"❌ Output formatting error: {str(e)}"
        })

    return updates


def create_docx_contract(contract_text: str, output_path: str, merged_data: Dict[str, Any]):
    """Create a formatted DOCX document from the contract text."""
    doc = Document()

    # Set document properties
    doc.core_properties.title = f"Subcontractor Agreement - {merged_data.get('project_name', 'Project')}"
    doc.core_properties.subject = "Nachunternehmervertrag / Subcontractor Agreement"
    doc.core_properties.author = "Contract Generation System"

    # Add title
    title = doc.add_heading('NACHUNTERNEHMERVERTRAG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('(SUBCONTRACTOR AGREEMENT)', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process contract text
    lines = contract_text.split('\n')
    current_paragraph = None

    for line in lines:
        # Skip the title lines we already added
        if line.strip() in ['NACHUNTERNEHMERVERTRAG', '(SUBCONTRACTOR AGREEMENT)']:
            continue

        # Handle section headers
        if line.startswith('§') or line.startswith('================'):
            if current_paragraph:
                current_paragraph = None
            if line.startswith('§'):
                doc.add_heading(line, 2)
            elif '=' in line:
                doc.add_page_break()
        elif line.startswith('---'):
            # Skip separator lines
            continue
        else:
            # Add regular text
            if line.strip():
                if line.startswith('    '):
                    # Indented text
                    p = doc.add_paragraph(line.strip())
                    p.paragraph_format.left_indent = Inches(0.5)
                else:
                    p = doc.add_paragraph(line)

                # Format specific elements
                if 'GESAMTSUMME' in line or 'TOTAL AMOUNT' in line:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(12)

    # Save document
    doc.save(output_path)