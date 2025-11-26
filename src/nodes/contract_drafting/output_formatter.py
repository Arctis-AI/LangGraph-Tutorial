"""Output formatter node for generating final contract files."""

import os
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.models.contract_drafting_state import ContractDraftingState


def output_formatter_node(state: ContractDraftingState) -> Dict[str, Any]:
    """
    Format and export contract.

    Generates:
    - TXT version
    - DOCX version with formatting
    - Quality report

    Args:
        state: Current workflow state

    Returns:
        Updates to state with output file paths
    """
    print("📄 Formatting output...")

    updates = {
        "current_step": "output_formatter",
        "messages": [],
        "output_files": {}
    }

    contract_draft = state.get("contract_draft", "")
    contract_type_data = state.get("contract_type_data", {})
    quality_report = state.get("quality_report", {})
    consistency_issues = state.get("consistency_issues", [])

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = "data/output"
    os.makedirs(output_dir, exist_ok=True)

    # 1. Save text version
    txt_filename = f"contract_{contract_type_data.get('code', 'draft')}_{timestamp}.txt"
    txt_path = os.path.join(output_dir, txt_filename)

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(contract_draft)

    updates["output_files"]["txt"] = txt_path

    # 2. Generate DOCX
    try:
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title
        title = doc.add_heading(contract_type_data.get('name_de', contract_type_data.get('name')), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"Vertragstyp: {contract_type_data.get('name')}\n").bold = True
        metadata_para.add_run(f"Code: {contract_type_data.get('code')}\n")
        metadata_para.add_run(f"Generiert: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
        metadata_para.add_run(f"Qualitätsscore: {quality_report.get('score', 0):.1f}/100")
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Add contract sections
        sections = contract_draft.split('\n\n')
        for para_text in sections:
            if not para_text.strip():
                continue

            # Check if it's a section header (starts with §)
            if para_text.strip().startswith('§'):
                # Section heading
                heading = doc.add_heading(para_text.strip(), level=1)
                heading_run = heading.runs[0]
                heading_run.font.size = Pt(14)
            elif para_text.strip().startswith('==='):
                # Separator - skip
                continue
            elif para_text.strip().startswith('---'):
                # Subsection separator - skip
                continue
            else:
                # Regular paragraph
                para = doc.add_paragraph(para_text.strip())
                para_run = para.runs[0] if para.runs else None
                if para_run:
                    para_run.font.size = Pt(11)

                # Highlight review markers
                if "[PRÜFUNG ERFORDERLICH:" in para_text or "[REVIEW NEEDED:" in para_text:
                    para.runs[0].font.highlight_color = 3  # Yellow highlight

        # Add quality summary page
        doc.add_page_break()
        doc.add_heading("Qualitätsbericht / Quality Report", level=1)

        quality_para = doc.add_paragraph()
        quality_para.add_run(f"Gesamtscore: {quality_report.get('score', 0):.1f}/100\n").bold = True
        quality_para.add_run(f"Bewertung: {quality_report.get('level', 'N/A')}\n\n")
        quality_para.add_run(f"Generierte Abschnitte: {quality_report.get('sections_generated', 0)}/{quality_report.get('sections_required', 0)}\n")
        quality_para.add_run(f"Vertragslänge: {quality_report.get('contract_length', 0)} Zeichen\n")
        quality_para.add_run(f"Datennutzung: {quality_report.get('data_usage_ratio', 0)*100:.0f}%\n")

        if consistency_issues:
            doc.add_heading("Gefundene Probleme / Issues Found", level=2)

            for issue in consistency_issues[:10]:  # Limit to first 10
                issue_para = doc.add_paragraph(style='List Bullet')
                severity = issue.get('severity', 'unknown')
                message = issue.get('message', 'Unknown issue')
                issue_para.add_run(f"[{severity.upper()}] ").bold = True
                issue_para.add_run(message)

        # Save DOCX
        docx_filename = f"contract_{contract_type_data.get('code', 'draft')}_{timestamp}.docx"
        docx_path = os.path.join(output_dir, docx_filename)
        doc.save(docx_path)

        updates["output_files"]["docx"] = docx_path

    except Exception as e:
        print(f"⚠️ Failed to generate DOCX: {e}")
        updates["messages"].append({
            "role": "system",
            "content": f"⚠️ DOCX generation failed: {str(e)}"
        })

    # 3. Save quality report as JSON
    import json
    report_filename = f"quality_report_{timestamp}.json"
    report_path = os.path.join(output_dir, report_filename)

    with open(report_path, "w", encoding="utf-8") as f:
        json.dump({
            "quality_report": quality_report,
            "consistency_issues": consistency_issues,
            "contract_type": contract_type_data.get("code"),
            "generated_at": datetime.now().isoformat()
        }, f, indent=2, ensure_ascii=False)

    updates["output_files"]["report"] = report_path
    updates["output_path"] = txt_path  # Primary output

    message = "✓ Generated outputs:\n"
    for file_type, path in updates["output_files"].items():
        message += f"  - {file_type.upper()}: {os.path.basename(path)}\n"

    updates["messages"].append({
        "role": "system",
        "content": message
    })

    updates["processing_status"] = "completed"

    return updates
