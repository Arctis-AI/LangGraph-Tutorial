"""Document extractor node for processing Verhandlungsprotokoll (supports DOCX, PDF, TXT)."""

import pdfplumber
from docx import Document
from typing import Dict, Any
from datetime import date, datetime
from src.core.llm_clients import get_llm_client
from src.models.state import ContractState
from src.models.contract import VerhandlungsprotokollData, ContractParty, PaymentTerms
from src.prompts import DOCUMENT_EXTRACTION_PROMPT, FIELD_EXTRACTION_PROMPT_TEMPLATE
import json


def document_extractor_node(state: ContractState) -> Dict[str, Any]:
    """
    Extract data from Verhandlungsprotokoll document (DOCX, PDF, or TXT) using LLM.
    """
    print("📄 Extracting data from Verhandlungsprotokoll...")

    updates = {
        "current_step": "document_extractor",
        "messages": []
    }

    doc_path = state.get("pdf_path")  # Key kept as pdf_path for compatibility
    if not doc_path:
        updates["messages"].append({
            "role": "system",
            "content": "⚠️ No document file available, skipping extraction"
        })
        return updates

    try:
        # Extract text from docx, PDF, or text file
        full_text = ""
        if doc_path.endswith('.txt'):
            # Read text file directly
            with open(doc_path, 'r', encoding='utf-8') as f:
                full_text = f.read()
        elif doc_path.endswith('.docx'):
            # Extract from Word document
            doc = Document(doc_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            full_text = "\n".join(paragraphs)
        else:
            # Extract from PDF
            with pdfplumber.open(doc_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        full_text += page_text + "\n"

        updates["verhandlungsprotokoll_raw"] = full_text
        updates["messages"].append({
            "role": "system",
            "content": f"✓ Extracted {len(full_text)} characters from document"
        })

        # Use LLM to structure the extracted data
        llm = get_llm_client()  # Uses default provider from config

        extraction_prompt = DOCUMENT_EXTRACTION_PROMPT(full_text)

        response = llm.invoke([
            {"role": "system", "content": "You are a precise data extraction assistant. Always return valid, complete JSON."},
            {"role": "user", "content": extraction_prompt}
        ])

        # Remove debug output for cleaner logs

        # Parse the LLM response as JSON
        try:
            # Clean up the response to extract JSON
            response_text = response.content

            # Remove markdown code block if present
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0]

            # Try to find JSON in the response (in case LLM added extra text)
            import re
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                response_text = json_match.group()

            extracted_data = json.loads(response_text)

            # Convert date strings to date objects where needed
            for date_field in ["negotiation_date", "contract_start_date", "contract_end_date"]:
                if extracted_data.get(date_field) and extracted_data[date_field] != "null":
                    try:
                        extracted_data[date_field] = datetime.strptime(
                            extracted_data[date_field], "%Y-%m-%d"
                        ).date()
                    except:
                        extracted_data[date_field] = None

            # Create structured objects with None handling (German fallbacks for consistency)
            contractor_data = extracted_data["contractor"]
            contractor_data["name"] = contractor_data.get("name") or "[Auftraggeber Name nicht gefunden]"
            contractor_data["address"] = contractor_data.get("address") or "[Auftraggeber Adresse nicht gefunden]"
            extracted_data["contractor"] = ContractParty(**contractor_data)

            subcontractor_data = extracted_data["subcontractor"]
            subcontractor_data["name"] = subcontractor_data.get("name") or "[Nachunternehmer Name nicht gefunden]"
            subcontractor_data["address"] = subcontractor_data.get("address") or "[Nachunternehmer Adresse nicht gefunden]"
            extracted_data["subcontractor"] = ContractParty(**subcontractor_data)

            payment_data = extracted_data["payment_terms"]
            payment_data["payment_schedule"] = payment_data.get("payment_schedule") or "Zahlungsbedingungen noch zu definieren"
            extracted_data["payment_terms"] = PaymentTerms(**payment_data)

            updates["verhandlungsprotokoll_data"] = extracted_data
            updates["messages"].append({
                "role": "system",
                "content": f"✓ Successfully extracted data from Verhandlungsprotokoll:\n" +
                         f"  Project: {extracted_data.get('project_name')}\n" +
                         f"  Contractor: {extracted_data['contractor'].name}\n" +
                         f"  Subcontractor: {extracted_data['subcontractor'].name}"
            })

        except (json.JSONDecodeError, KeyError, TypeError) as e:
            # Log the actual error for debugging
            updates["messages"].append({
                "role": "system",
                "content": f"⚠️ Error parsing LLM response: {str(e)}\nResponse was: {response.content[:500]}"
            })
            # Try a simpler extraction as fallback
            updates["verhandlungsprotokoll_data"] = extract_with_fallback(full_text)
            updates["messages"].append({
                "role": "system",
                "content": "⚠️ Using simplified extraction method"
            })

    except Exception as e:
        updates["error"] = f"PDF extraction failed: {str(e)}"
        updates["messages"].append({
            "role": "system",
            "content": f"❌ PDF extraction error: {str(e)}"
        })
        # Don't use defaults - let it fail or extract what it can
        updates["verhandlungsprotokoll_data"] = None

    return updates


def extract_with_fallback(text: str) -> Dict[str, Any]:
    """Simpler extraction method using pattern matching and LLM for specific fields."""
    import re

    try:
        llm = get_llm_client()  # Uses default provider from config

        # Extract specific fields with targeted prompts
        def extract_field(field_name: str, prompt: str) -> str:
            field_prompt = FIELD_EXTRACTION_PROMPT_TEMPLATE(field_name, prompt, text)
            response = llm.invoke([
                {"role": "system", "content": "Extract only the requested information. Be concise."},
                {"role": "user", "content": field_prompt}
            ])
            return response.content.strip()

        # Extract key information with targeted prompts
        project_name = extract_field("project_name",
            "What is the project name? Return ONLY the project name, nothing else.")

        project_location = extract_field("project_location",
            "What is the project location/address? Return ONLY the location.")

        # Extract contractor info
        contractor_info = extract_field("contractor",
            "Who is the main contractor (Auftraggeber)? Return name and address only.")
        contractor_parts = contractor_info.split('\n')
        contractor_name = contractor_parts[0] if contractor_parts else ""
        contractor_address = contractor_parts[1] if len(contractor_parts) > 1 else ""

        # Extract subcontractor info
        subcontractor_info = extract_field("subcontractor",
            "Who is the subcontractor (Nachunternehmer)? Return name and address only.")
        subcontractor_parts = subcontractor_info.split('\n')
        subcontractor_name = subcontractor_parts[0] if subcontractor_parts else ""
        subcontractor_address = subcontractor_parts[1] if len(subcontractor_parts) > 1 else ""

        # Try to find dates with regex
        date_pattern = r'(\d{1,2})\.(\d{1,2})\.(\d{4})'
        dates = re.findall(date_pattern, text)

        # Use first found date as start, calculate end as 6 months later
        if dates:
            day, month, year = dates[0]
            start_date = date(int(year), int(month), int(day))
        else:
            start_date = date.today()

        end_date = date(start_date.year + (1 if start_date.month > 6 else 0),
                        (start_date.month + 6) % 12 or 12, start_date.day)

        return {
            "project_name": project_name,
            "project_location": project_location,
            "project_description": extract_field("description",
                "Briefly describe the project scope. Maximum 2 sentences."),
            "contractor": ContractParty(
                name=contractor_name or "[Auftraggeber Name nicht gefunden]",
                address=contractor_address or "[Auftraggeber Adresse nicht gefunden]"
            ),
            "subcontractor": ContractParty(
                name=subcontractor_name or "[Nachunternehmer Name nicht gefunden]",
                address=subcontractor_address or "[Nachunternehmer Adresse nicht gefunden]"
            ),
            "contract_start_date": start_date,
            "contract_end_date": end_date,
            "scope_of_work": extract_field("scope",
                "What work will be performed? Summarize in 2-3 sentences."),
            "payment_terms": PaymentTerms(
                payment_schedule=extract_field("payment",
                    "What are the payment terms? Return in one sentence.") or "Zahlungsbedingungen noch zu definieren",
                payment_deadline_days=30
            ),
            "special_agreements": [],
            "excluded_services": []
        }
    except Exception as e:
        # If extraction completely fails, return None
        return None


